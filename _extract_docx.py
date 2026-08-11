# -*- coding: utf-8 -*-
import docx
d = docx.Document(r"D:\work\林育丞-ai开发.docx")
lines = []
for p in d.paragraphs:
    t = p.text.strip()
    if t:
        lines.append(t)
for tbl in d.tables:
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        line = " | ".join([c for c in cells if c])
        if line:
            lines.append(line)
text = "\n".join(lines)
out = r"E:\job3.0\_resume_docx.txt"
with open(out, "w", encoding="utf-8") as f:
    f.write(text)
print("chars:", len(text))
print("first 600 chars:\n", text[:600])
